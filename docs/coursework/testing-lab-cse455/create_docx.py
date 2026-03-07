"""Generate DOCX files for CSE-455 Software Testing Lab assignments."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement

SCREENSHOTS_DIR = os.path.join(os.path.dirname(__file__), "screenshots")
OUTPUT_DIR = os.path.dirname(__file__)


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '1A73E8')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def add_url_section(doc, title, urls):
    """Add a references section with clickable URLs."""
    doc.add_page_break()
    doc.add_heading(title, level=1)
    for label, url in urls:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        add_hyperlink(p, url, url)


def add_cover_info(doc, title, subtitle, student_name, student_id, course, instructor, date):
    """Add a simple cover heading block."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    info_lines = [
        f"Student: {student_name}",
        f"Student ID: {student_id}",
        f"Course: {course}",
        f"Instructor: {instructor}",
        f"Date: {date}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.add_page_break()


def add_screenshot(doc, image_path, caption, width=Inches(6.2)):
    """Add an image with a caption below it."""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    else:
        doc.add_paragraph(f"[Image not found: {image_path}]")
    doc.add_paragraph()


# ============================================================
# ASSIGNMENT 1: Theory Assignment - Test Cases
# ============================================================
def create_assignment1():
    doc = Document()

    # Adjust default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover_info(
        doc,
        title="Theory Assignment 1",
        subtitle="FYP Test Cases — Jira Test Management",
        student_name="Ghulam Mujtaba",
        student_id="FA21-BSE-084",
        course="CSE-455 Software Testing",
        instructor="Dr. Instructor",
        date="February 2026",
    )

    # Section 1
    doc.add_heading("1. Test Cases Overview", level=1)
    doc.add_paragraph(
        "21 test cases were created in Jira for the MegiLance FYP project, "
        "covering authentication, project management, proposals, contracts, "
        "payments, profiles, and AI matching modules."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "01-test-cases-list.png"),
        "Figure 1: Jira Test Cases List — 21 of 21 test cases (label: test-case)",
    )

    # Section 2
    doc.add_heading("2. Test Case — Project Management", level=1)
    doc.add_paragraph(
        "Each test case includes: Test Case ID, description, preconditions, "
        "test steps, expected results, priority, and labels (black-box, functional, test-case)."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "02-test-case-detail.png"),
        "Figure 2: MEG-105 — TC-PROJ-001: Create New Project with Valid Data",
    )

    # Section 3
    doc.add_heading("3. Test Case — Profile Management", level=1)
    doc.add_paragraph(
        "Profile management test case covering freelancer profile update with valid data."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "08-test-case-MEG106.png"),
        "Figure 3: MEG-106 — TC-PROF-001: Update Freelancer Profile with Valid Data",
    )

    # Section 4
    doc.add_heading("4. Test Case — Proposal Management", level=1)
    doc.add_paragraph(
        "Proposal management test case covering submission of valid proposals on open projects."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "09-test-case-MEG107.png"),
        "Figure 4: MEG-107 — TC-PROP-001: Submit Valid Proposal on Open Project",
    )

    # Section 5
    doc.add_heading("5. Test Case — Payment System", level=1)
    doc.add_paragraph(
        "Payment system test case covering milestone payment processing."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "10-test-case-MEG108.png"),
        "Figure 5: MEG-108 — TC-PAY-001: Process Valid Milestone Payment",
    )

    # Section 6
    doc.add_heading("6. Test Case — Search & AI Matching", level=1)
    doc.add_paragraph(
        "Search and AI matching test case covering freelancer skill-based search."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "11-test-case-MEG109.png"),
        "Figure 6: MEG-109 — TC-SEARCH-001: Search Freelancers by Skill Keyword",
    )

    # Section 7
    doc.add_heading("7. Test Case Categories", level=1)
    doc.add_paragraph(
        "Test cases span multiple modules:"
    )
    bullets = [
        "Authentication & Security (login, registration, JWT, password reset)",
        "Project Management (create, browse, search, filter)",
        "Proposals & Contracts (submit, accept, milestone tracking)",
        "Payments & Invoices (escrow, withdrawals, refunds)",
        "User Profiles (freelancer/client profiles, portfolio, reviews)",
        "AI Matching Engine (skill matching, recommendations)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # Section 8: Zephyr Scale Test Cases
    doc.add_page_break()
    doc.add_heading("8. Zephyr Scale — Test Case Management", level=1)
    doc.add_paragraph(
        "In addition to Jira issues, 10 test cases were created using Zephyr Scale "
        "(SmartBear Zephyr Advanced), a dedicated test management tool integrated with Jira. "
        "Zephyr Scale provides structured test case management with test steps, "
        "test cycles, test plans, and execution tracking."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-test-cases-list-all.png"),
        "Figure 7: Zephyr Scale — All 10 Test Cases (MEG-T1 to MEG-T10)",
    )

    doc.add_heading("8.1 Zephyr Scale Test Cases List", level=2)
    zephyr_cases = [
        ("MEG-T1", "TC-AUTH-001: User Registration with Valid Data"),
        ("MEG-T2", "TC-AUTH-002: User Login with Valid Credentials"),
        ("MEG-T3", "TC-AUTH-003: User Login with Invalid Password"),
        ("MEG-T4", "TC-AUTH-004: Password Reset Request"),
        ("MEG-T5", "TC-AUTH-005: JWT Token Refresh"),
        ("MEG-T6", "TC-PROJ-001: Create New Project"),
        ("MEG-T7", "TC-PROF-001: Update User Profile"),
        ("MEG-T8", "TC-PROP-001: Submit Proposal for Project"),
        ("MEG-T9", "TC-PAY-001: Process Payment Transaction"),
        ("MEG-T10", "TC-SEARCH-001: Search Projects with Filters"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Key"
    hdr[1].text = "Test Case Name"
    hdr[2].text = "Status"
    for key, name in zephyr_cases:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = name
        row[2].text = "Draft"
    doc.add_paragraph()

    doc.add_heading("8.2 Individual Test Case — Details View", level=2)
    doc.add_paragraph(
        "Each Zephyr Scale test case includes: Name, Objective, Precondition, "
        "Status, Priority, Owner, and a Test Script tab for step-by-step test steps."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-tc1-details.png"),
        "Figure 8: MEG-T1 — TC-AUTH-001: User Registration with Valid Data (Details View)",
    )

    # Jira & Zephyr URLs Reference
    add_url_section(doc, "9. Jira & Zephyr URLs Reference", [
        ("All 21 Test Cases — Jira Issues",
         "https://ghulam-mujtaba.atlassian.net/issues/?jql=project%20%3D%20MEG%20AND%20labels%20%3D%20test-case"),
        ("TC-PROJ-001 — Create New Project (MEG-105)",
         "https://ghulam-mujtaba.atlassian.net/browse/MEG-105"),
        ("TC-PROF-001 — Update Freelancer Profile (MEG-106)",
         "https://ghulam-mujtaba.atlassian.net/browse/MEG-106"),
        ("TC-PROP-001 — Submit Valid Proposal (MEG-107)",
         "https://ghulam-mujtaba.atlassian.net/browse/MEG-107"),
        ("TC-PAY-001 — Process Milestone Payment (MEG-108)",
         "https://ghulam-mujtaba.atlassian.net/browse/MEG-108"),
        ("TC-SEARCH-001 — Search by Skill Keyword (MEG-109)",
         "https://ghulam-mujtaba.atlassian.net/browse/MEG-109"),
        ("Zephyr Scale — All Test Cases",
         "https://ghulam-mujtaba.atlassian.net/projects/MEG?selectedItem=com.atlassian.plugins.atlassian-connect-plugin:com.kanoah.test-manager__main-project-page"),
    ])

    output_path = os.path.join(OUTPUT_DIR, "CSE455_Assignment1_TestCases_GhulamMujtaba.docx")
    doc.save(output_path)
    print(f"Assignment 1 DOCX saved: {output_path}")


# ============================================================
# ASSIGNMENT 2: Lab Assignment - Kanban/Scrum & Automation
# ============================================================
def create_assignment2():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover_info(
        doc,
        title="Lab Assignment 2",
        subtitle="Kanban/Scrum Board & Jira Automation Rules",
        student_name="Ghulam Mujtaba",
        student_id="FA21-BSE-084",
        course="CSE-455 Software Testing",
        instructor="Dr. Instructor",
        date="February 2026",
    )

    # Section 1
    doc.add_heading("1. Scrum Board", level=1)
    doc.add_paragraph(
        "The MegiLance project uses a Scrum board with TO DO, IN PROGRESS, and DONE columns."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "03-scrum-board.png"),
        "Figure 1: MEG Scrum Board — TO DO / IN PROGRESS / DONE columns",
    )

    # Section 2
    doc.add_heading("2. Backlog & Sprints", level=1)
    doc.add_paragraph(
        "Sprints are organized in the backlog with defined start/end dates and assigned work items."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "04-backlog-sprints.png"),
        "Figure 2: Backlog — Sprint planning with MEG Sprint 1 and Foundation Sprint",
    )

    # Section 3
    doc.add_heading("3. Timeline & Epics", level=1)
    doc.add_paragraph(
        "Epics provide high-level feature grouping across the project timeline."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "05-timeline-epics.png"),
        "Figure 3: Timeline view — Epics (Auth, AI Matching, Client/Freelancer/Admin Portals)",
    )

    # Section 4
    doc.add_heading("4. Issue Types (Epics, Stories, Bugs, Tasks)", level=1)
    doc.add_paragraph(
        "The project tracks 88 issues across multiple types: Epics, Stories, Bugs, and Tasks."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "07-all-issues-types.png"),
        "Figure 4: All work items — 88 issues filtered by type (Epic, Story, Bug, Task)",
    )

    # Section 5
    doc.add_heading("5. Automation Rules", level=1)
    doc.add_paragraph(
        "4 Jira automation rules are active to streamline workflow:"
    )
    rules = [
        "When a bug is created → set due date based on priority",
        "When all sub-tasks are done → move parent to done",
        "When a task is near due → send email reminder",
        "When a work item is transitioned → automatically assign",
    ]
    for r in rules:
        doc.add_paragraph(r, style="List Bullet")

    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "06-automation-rules.png"),
        "Figure 5: All 4 automation rules — enabled and active",
    )

    # Jira URLs Reference
    add_url_section(doc, "6. Jira URLs Reference", [
        ("Scrum Board (TO DO / IN PROGRESS / DONE)",
         "https://ghulam-mujtaba.atlassian.net/jira/software/c/projects/MEG/boards/3"),
        ("Backlog & Sprints",
         "https://ghulam-mujtaba.atlassian.net/jira/software/c/projects/MEG/boards/3/backlog"),
        ("Timeline / Epics",
         "https://ghulam-mujtaba.atlassian.net/jira/software/c/projects/MEG/boards/3/timeline"),
        ("All Issues (Epic, Story, Bug, Task)",
         "https://ghulam-mujtaba.atlassian.net/issues/?jql=project%20%3D%20MEG%20AND%20type%20in%20(Epic%2C%20Story%2C%20Bug%2C%20Task)%20ORDER%20BY%20type%20ASC"),
        ("Automation Rules",
         "https://ghulam-mujtaba.atlassian.net/jira/software/c/projects/MEG/settings/automate"),
    ])

    output_path = os.path.join(OUTPUT_DIR, "CSE455_Assignment2_KanbanAutomation_GhulamMujtaba.docx")
    doc.save(output_path)
    print(f"Assignment 2 DOCX saved: {output_path}")


# ============================================================
# ZEPHYR SCALE ASSIGNMENT: Test Case Management in Zephyr Scale
# ============================================================
def _add_test_case_table(doc, tc_id, name, objective, preconditions, steps, priority="Normal"):
    """Add a formatted table for a single Zephyr Scale test case."""
    doc.add_heading(f"{tc_id}: {name}", level=2)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"

    info_rows = [
        ("Test Case ID", tc_id),
        ("Name", name),
        ("Objective", objective),
        ("Preconditions", preconditions),
        ("Priority", priority),
        ("Status", "Draft"),
        ("Version", "1.0"),
    ]
    for label, value in info_rows:
        row = tbl.add_row().cells
        row[0].text = label
        run = row[0].paragraphs[0].runs[0]
        run.bold = True
        row[1].text = value

    doc.add_paragraph()
    # Steps sub-table
    step_tbl = doc.add_table(rows=1, cols=4)
    step_tbl.style = "Light Grid Accent 1"
    hdr = step_tbl.rows[0].cells
    hdr[0].text = "Step #"
    hdr[1].text = "Action"
    hdr[2].text = "Test Data"
    hdr[3].text = "Expected Result"
    for step in steps:
        row = step_tbl.add_row().cells
        row[0].text = str(step[0])
        row[1].text = step[1]
        row[2].text = step[2]
        row[3].text = step[3]
    doc.add_paragraph()


def create_zephyr_assignment():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover_info(
        doc,
        title="Theory Assignment 1 — Zephyr Scale",
        subtitle="Test Case Management Using Zephyr Scale (Jira Plugin)\nMegiLance — AI-Powered Freelancing Platform",
        student_name="Ghulam Mujtaba",
        student_id="FA21-BSE-084",
        course="CSE-455 Software Testing",
        instructor="Dr. Instructor",
        date="March 2026",
    )

    # ── Table of Contents ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Zephyr Scale Installation & Setup",
        "   2.1 Atlassian Marketplace",
        "   2.2 Verifying Installation in Connected Apps",
        "3. Zephyr Scale — Test Cases Overview",
        "   3.1 All Test Cases List (10 Test Cases)",
        "   3.2 Test Cases Summary Table",
        "4. Detailed Test Cases with Screenshots",
        "   4.1 MEG-T1: User Registration (Details View)",
        "   4.2 MEG-T1: User Registration (Test Script View)",
        "   4.3 MEG-T6: Create New Project (Details View)",
        "   4.4 MEG-T8: Submit Proposal (Details View)",
        "   4.5 MEG-T9: Process Payment (Details View)",
        "   4.6 MEG-T10: Search Projects (Details View)",
        "5. Complete Test Case Specifications",
        "6. Zephyr Scale — Test Cycles",
        "7. Zephyr Scale — Test Plans",
        "8. Zephyr Scale — Reports & Dashboards",
        "9. Zephyr Scale Features Used",
        "10. Jira Issues vs Zephyr Scale — Comparison",
        "11. Jira Project — Supporting Evidence",
        "   11.1 Scrum Board",
        "   11.2 Backlog & Sprint Planning",
        "   11.3 Timeline / Epics",
        "   11.4 All Work Items (88 Issues)",
        "   11.5 Automation Rules",
        "   11.6 Jira Test Case Issues (21 Issues)",
        "12. Conclusion",
        "13. URLs Reference",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        if not item.startswith("   "):
            p.runs[0].bold = True
    doc.add_page_break()

    # ── Section 1: Introduction ──
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This assignment demonstrates the use of Zephyr Scale (SmartBear Zephyr Advanced), "
        "a professional test management tool integrated with Jira Cloud, for creating and "
        "managing structured test cases for the MegiLance FYP project — an AI-powered "
        "freelancing platform built with Next.js 16, React 19, and FastAPI."
    )
    doc.add_paragraph(
        "Zephyr Scale provides dedicated test case management capabilities including "
        "structured test steps, test cycles, test plans, execution tracking, and reporting — "
        "going beyond basic Jira issues to offer a complete test management solution."
    )
    doc.add_paragraph(
        "The MegiLance platform connects clients with freelancers through intelligent "
        "AI-powered matching, managing the full project lifecycle from posting to payment. "
        "The system includes 88+ Jira work items across Epics, Stories, Bugs, and Tasks, "
        "with 21 dedicated test case issues and 10 Zephyr Scale structured test cases."
    )

    # ── Section 2: Installation ──
    doc.add_page_break()
    doc.add_heading("2. Zephyr Scale Installation & Setup", level=1)

    doc.add_heading("2.1 Atlassian Marketplace", level=2)
    doc.add_paragraph(
        "Zephyr Scale (Zephyr — Test Management and Automation for Jira) was installed "
        "from the Atlassian Marketplace as a free trial and enabled for the MegiLance (MEG) "
        "project. The plugin provides enterprise-grade test management integrated directly "
        "into the Jira Cloud interface."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-install.png"),
        "Figure 1: Zephyr Scale — Atlassian Marketplace Overview Page (SmartBear)",
    )

    doc.add_heading("2.2 Verifying Installation in Connected Apps", level=2)
    doc.add_paragraph(
        "After installation, Zephyr Scale appears in the Jira Administration panel under "
        "Connected Apps / Manage Apps, confirming 'Zephyr — Test Management and Automation "
        "for Jira' by SmartBear is active with the Advanced edition."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-check.png"),
        "Figure 2: Jira Admin — Connected Apps showing Zephyr Scale installed",
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-admin-apps.png"),
        "Figure 3: Jira Administration — Manage Apps (Zephyr Scale Active)",
    )

    # ── Section 3: Test Cases Overview ──
    doc.add_page_break()
    doc.add_heading("3. Zephyr Scale — Test Cases Overview", level=1)

    doc.add_heading("3.1 All Test Cases List (10 Test Cases)", level=2)
    doc.add_paragraph(
        "10 test cases were created in Zephyr Scale covering the core modules of the "
        "MegiLance platform. The test cases are accessed through the Zephyr tab in the "
        "MegiLance project navigation, which shows the SmartBear Zephyr Advanced interface "
        "with Test Cases, Test Cycles, Test Plans, and Reports tabs."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-test-cases-list-all.png"),
        "Figure 4: Zephyr Scale — All 10 Test Cases (MEG-T1 to MEG-T10) with Priority, Version, Name, and Status columns",
    )

    doc.add_heading("3.2 Test Cases Summary Table", level=2)
    doc.add_paragraph(
        "The following table summarizes all 10 test cases organized by module:"
    )

    summary_table = doc.add_table(rows=1, cols=5)
    summary_table.style = "Light Grid Accent 1"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Key"
    hdr[1].text = "Test Case Name"
    hdr[2].text = "Module"
    hdr[3].text = "Priority"
    hdr[4].text = "Status"
    summary_data = [
        ("MEG-T1", "TC-AUTH-001: User Registration with Valid Data", "Authentication", "Normal"),
        ("MEG-T2", "TC-AUTH-002: User Login with Valid Credentials", "Authentication", "Normal"),
        ("MEG-T3", "TC-AUTH-003: User Login with Invalid Password", "Authentication", "Normal"),
        ("MEG-T4", "TC-AUTH-004: Password Reset Request", "Authentication", "Normal"),
        ("MEG-T5", "TC-AUTH-005: JWT Token Refresh", "Authentication", "Normal"),
        ("MEG-T6", "TC-PROJ-001: Create New Project", "Project Mgmt", "Normal"),
        ("MEG-T7", "TC-PROF-001: Update User Profile", "User Profiles", "Normal"),
        ("MEG-T8", "TC-PROP-001: Submit Proposal for Project", "Proposals", "Normal"),
        ("MEG-T9", "TC-PAY-001: Process Payment Transaction", "Payments", "Normal"),
        ("MEG-T10", "TC-SEARCH-001: Search Projects with Filters", "Search", "Normal"),
    ]
    for key, name, module, priority in summary_data:
        row = summary_table.add_row().cells
        row[0].text = key
        row[1].text = name
        row[2].text = module
        row[3].text = priority
        row[4].text = "Draft"
    doc.add_paragraph()

    # ── Section 4: Detailed Test Cases with Screenshots ──
    doc.add_page_break()
    doc.add_heading("4. Detailed Test Cases with Screenshots", level=1)
    doc.add_paragraph(
        "This section shows the actual Zephyr Scale interface for each test case, "
        "including the Details view (Name, Objective, Precondition, Status, Priority, Owner) "
        "and the Test Script tab (step-by-step test steps with Test Data and Expected Results)."
    )

    doc.add_heading("4.1 MEG-T1: User Registration — Details View", level=2)
    doc.add_paragraph(
        "The Details view for MEG-T1 shows the test case metadata including Name, "
        "Objective, Precondition, Status (Draft), Priority (Normal), and Owner assignment. "
        "This is the primary view for reviewing test case information in Zephyr Scale."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-tc1-details.png"),
        "Figure 5: MEG-T1 — TC-AUTH-001: User Registration with Valid Data (Zephyr Scale Details View)",
    )

    doc.add_heading("4.2 MEG-T1: User Registration — Test Script View", level=2)
    doc.add_paragraph(
        "The Test Script tab provides a structured step-by-step test execution guide. "
        "Each step includes an Action description, Test Data, and Expected Result. "
        "This structured format ensures consistent test execution across team members."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-tc1-testscript.png"),
        "Figure 6: MEG-T1 — TC-AUTH-001: Test Script Tab (Step-by-Step Test Steps)",
    )

    doc.add_heading("4.3 MEG-T6: Create New Project — Details View", level=2)
    doc.add_paragraph(
        "MEG-T6 covers the Project Management module, testing that a client user "
        "can create a new project with all required fields (title, description, budget, "
        "skills, deadline) and receive a successful HTTP 201 response."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-tc6-details.png"),
        "Figure 7: MEG-T6 — TC-PROJ-001: Create New Project (Zephyr Scale Details View)",
    )

    doc.add_heading("4.4 MEG-T8: Submit Proposal — Details View", level=2)
    doc.add_paragraph(
        "MEG-T8 covers the Proposals module, verifying that a freelancer can submit "
        "a proposal for an open project with proposed amount, cover letter, and "
        "estimated delivery time."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-tc8-details.png"),
        "Figure 8: MEG-T8 — TC-PROP-001: Submit Proposal for Project (Zephyr Scale Details View)",
    )

    doc.add_heading("4.5 MEG-T9: Process Payment — Details View", level=2)
    doc.add_paragraph(
        "MEG-T9 covers the Payments module, testing that a payment can be processed "
        "for an accepted contract milestone with proper balance updates and "
        "transaction record creation."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-tc9-details.png"),
        "Figure 9: MEG-T9 — TC-PAY-001: Process Payment Transaction (Zephyr Scale Details View)",
    )

    doc.add_heading("4.6 MEG-T10: Search Projects — Details View", level=2)
    doc.add_paragraph(
        "MEG-T10 covers the Search module, verifying that users can search and filter "
        "projects using multiple criteria including keyword, skill filters, budget range, "
        "and sorting options."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-tc10-details.png"),
        "Figure 10: MEG-T10 — TC-SEARCH-001: Search Projects with Filters (Zephyr Scale Details View)",
    )

    # ── Section 5: Complete Test Case Specifications ──
    doc.add_page_break()
    doc.add_heading("5. Complete Test Case Specifications", level=1)
    doc.add_paragraph(
        "Each test case below includes the full specification: Objective, Preconditions, "
        "Priority, and a step-by-step Test Script with Test Data and Expected Results. "
        "These specifications match exactly what is configured in Zephyr Scale."
    )

    # TC-1: User Registration
    _add_test_case_table(doc, "MEG-T1", "TC-AUTH-001: User Registration with Valid Data",
        objective="Verify that a new user can register with valid credentials and receive a confirmation.",
        preconditions="Application is running; /signup page is accessible; no existing account with the test email.",
        steps=[
            (1, "Navigate to the registration page at /signup", "", "Registration form is displayed with email, password, full name, and role fields"),
            (2, "Enter a valid full name", "John Doe", "Name field accepts the input"),
            (3, "Enter a valid email address", "testuser@example.com", "Email field accepts the input"),
            (4, "Enter a strong password", "SecureP@ss123!", "Password field accepts input; strength indicator shows 'Strong'"),
            (5, "Select user role", "Freelancer", "Role dropdown shows selected value"),
            (6, "Click the 'Sign Up' button", "", "Loading spinner appears; form is submitted"),
            (7, "Verify success response", "", "HTTP 201; user object returned with id, email, role; JWT access and refresh tokens set"),
        ],
    )

    # TC-2: User Login Valid
    _add_test_case_table(doc, "MEG-T2", "TC-AUTH-002: User Login with Valid Credentials",
        objective="Verify that a registered user can log in with correct email and password.",
        preconditions="A user account exists with known credentials; /login page is accessible.",
        steps=[
            (1, "Navigate to /login page", "", "Login form displayed with email and password fields"),
            (2, "Enter registered email", "testuser@example.com", "Email field accepts input"),
            (3, "Enter correct password", "SecureP@ss123!", "Password field accepts input (masked)"),
            (4, "Click 'Log In' button", "", "Loading spinner appears; credentials submitted to POST /api/auth/login"),
            (5, "Verify login success", "", "HTTP 200; JWT access token (30 min) and refresh token (7 days) returned"),
            (6, "Verify redirect to dashboard", "", "User is redirected to /portal/dashboard with correct role"),
        ],
    )

    # TC-3: User Login Invalid Password
    _add_test_case_table(doc, "MEG-T3", "TC-AUTH-003: User Login with Invalid Password",
        objective="Verify that login is rejected when an incorrect password is provided.",
        preconditions="A user account exists with known email; /login page is accessible.",
        steps=[
            (1, "Navigate to /login page", "", "Login form displayed"),
            (2, "Enter registered email", "testuser@example.com", "Email field accepts input"),
            (3, "Enter incorrect password", "WrongPassword123", "Password field accepts input"),
            (4, "Click 'Log In' button", "", "Credentials submitted"),
            (5, "Verify error response", "", "HTTP 401; error message: 'Invalid credentials'"),
            (6, "Verify user stays on login page", "", "No redirect; login form remains visible with error message"),
        ],
    )

    # TC-4: Password Reset
    _add_test_case_table(doc, "MEG-T4", "TC-AUTH-004: Password Reset Request",
        objective="Verify that a user can request a password reset link via their registered email.",
        preconditions="A user account exists; /forgot-password page is accessible; SMTP service is configured.",
        steps=[
            (1, "Navigate to /forgot-password page", "", "Password reset form displayed with email field"),
            (2, "Enter the registered email", "testuser@example.com", "Email field accepts input"),
            (3, "Click 'Send Reset Link' button", "", "Request sent to POST /api/auth/forgot-password"),
            (4, "Verify success message", "", "Success message displayed: 'If the email exists, a reset link has been sent'"),
            (5, "Check email inbox", "", "Password reset email received with a valid reset link containing a token"),
        ],
    )

    # TC-5: JWT Token Refresh
    _add_test_case_table(doc, "MEG-T5", "TC-AUTH-005: JWT Token Refresh",
        objective="Verify that an expired access token can be refreshed using a valid refresh token.",
        preconditions="User is logged in with valid access and refresh tokens; access token is about to expire or expired.",
        steps=[
            (1, "Make an API request with an expired access token", "GET /api/auth/me with expired JWT", "HTTP 401 Unauthorized response"),
            (2, "Send refresh token to /api/auth/refresh", "POST with refresh_token cookie", "HTTP 200; new access token returned"),
            (3, "Retry original request with new access token", "GET /api/auth/me with new JWT", "HTTP 200; user profile data returned successfully"),
            (4, "Verify new token expiry", "", "New access token has 30-minute expiry; refresh token unchanged"),
        ],
    )

    # TC-6: Create New Project
    _add_test_case_table(doc, "MEG-T6", "TC-PROJ-001: Create New Project",
        objective="Verify that a client user can create a new project with all required fields.",
        preconditions="User is logged in as a Client; /portal/projects/create page is accessible.",
        steps=[
            (1, "Navigate to /portal/projects/create", "", "Project creation form is displayed"),
            (2, "Enter project title", "E-Commerce Website Redesign", "Title field accepts input (max 200 chars)"),
            (3, "Enter project description", "Full redesign of online store with React and Node.js backend", "Description field accepts input"),
            (4, "Set budget range", "Min: $500, Max: $2000", "Budget fields accept numeric input"),
            (5, "Select required skills", "React, Node.js, PostgreSQL", "Skills tags are added"),
            (6, "Set project deadline", "2026-06-30", "Date picker accepts future date"),
            (7, "Click 'Post Project' button", "", "HTTP 201; project created with status 'Open'; redirect to project detail page"),
        ],
    )

    # TC-7: Update User Profile
    _add_test_case_table(doc, "MEG-T7", "TC-PROF-001: Update User Profile",
        objective="Verify that a user can update their profile information successfully.",
        preconditions="User is logged in; /portal/profile/edit page is accessible.",
        steps=[
            (1, "Navigate to /portal/profile/edit", "", "Profile edit form is displayed with current data pre-filled"),
            (2, "Update full name", "John M. Doe", "Name field accepts the updated value"),
            (3, "Update bio/description", "Senior Full-Stack Developer with 5+ years experience", "Bio field accepts input (max 1000 chars)"),
            (4, "Upload new avatar image", "avatar.jpg (under 2 MB)", "Image preview shown; file accepted"),
            (5, "Add/update skills", "Python, FastAPI, Next.js, Docker", "Skills tags updated"),
            (6, "Click 'Save Profile' button", "", "HTTP 200; profile updated; success toast notification displayed"),
            (7, "Verify changes persisted", "", "Reload profile page; all updated fields show new values"),
        ],
    )

    # TC-8: Submit Proposal
    _add_test_case_table(doc, "MEG-T8", "TC-PROP-001: Submit Proposal for Project",
        objective="Verify that a freelancer can submit a proposal for an open project.",
        preconditions="User is logged in as Freelancer; an open project exists; user hasn't already submitted a proposal for this project.",
        steps=[
            (1, "Navigate to an open project detail page", "", "Project details displayed with 'Submit Proposal' button visible"),
            (2, "Click 'Submit Proposal' button", "", "Proposal form/modal is displayed"),
            (3, "Enter proposed amount", "$1200", "Amount field accepts numeric input within project budget"),
            (4, "Enter cover letter", "I have 5 years of experience in React and Node.js...", "Cover letter field accepts text input"),
            (5, "Set estimated delivery time", "30 days", "Duration field accepts input"),
            (6, "Click 'Submit' button", "", "HTTP 201; proposal created with status 'Pending'; confirmation message shown"),
            (7, "Verify proposal appears in user's proposals list", "", "Proposal visible in /portal/proposals with correct status"),
        ],
    )

    # TC-9: Process Payment
    _add_test_case_table(doc, "MEG-T9", "TC-PAY-001: Process Payment Transaction",
        objective="Verify that a payment can be processed for an accepted contract milestone.",
        preconditions="A contract exists with an accepted milestone; client has sufficient balance; payment gateway is configured.",
        steps=[
            (1, "Navigate to the contract detail page", "", "Contract details displayed with milestones list"),
            (2, "Click 'Release Payment' on a completed milestone", "", "Payment confirmation dialog appears"),
            (3, "Confirm the payment amount", "$400 (milestone amount)", "Amount displayed correctly"),
            (4, "Click 'Confirm Payment' button", "", "Payment processed; HTTP 200 response"),
            (5, "Verify payment status updated", "", "Milestone status changes to 'Paid'; transaction record created"),
            (6, "Verify freelancer balance updated", "", "Freelancer's account balance increased by milestone amount minus platform fee"),
        ],
    )

    # TC-10: Search Projects
    _add_test_case_table(doc, "MEG-T10", "TC-SEARCH-001: Search Projects with Filters",
        objective="Verify that users can search and filter projects using multiple criteria.",
        preconditions="Multiple projects exist in the system with varying skills, budgets, and categories.",
        steps=[
            (1, "Navigate to /projects (project listing page)", "", "Project listing page displayed with search bar and filter panel"),
            (2, "Enter search keyword", "React", "Search field accepts input"),
            (3, "Apply skill filter", "Select: React, Node.js", "Filter tags applied; results update dynamically"),
            (4, "Apply budget range filter", "Min: $500, Max: $2000", "Budget filter applied"),
            (5, "Apply sort option", "Sort by: Newest First", "Results re-ordered by creation date descending"),
            (6, "Verify filtered results", "", "Only projects matching all criteria are displayed; result count updated"),
            (7, "Click on a project from results", "", "Redirected to project detail page with correct data"),
        ],
    )

    # ── Section 6: Test Cycles ──
    doc.add_page_break()
    doc.add_heading("6. Zephyr Scale — Test Cycles", level=1)
    doc.add_paragraph(
        "Test Cycles in Zephyr Scale allow grouping test cases into execution sets "
        "for sprint-based or release-based testing. Test cycles track which test cases "
        "have been executed, their pass/fail status, and execution history."
    )
    doc.add_paragraph(
        "The MegiLance project's Test Cycles view is accessible from the Zephyr Scale "
        "navigation tab and shows all configured test execution cycles."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-test-cycles.png"),
        "Figure 11: Zephyr Scale — Test Cycles View (Test Execution Cycles)",
    )

    # ── Section 7: Test Plans ──
    doc.add_page_break()
    doc.add_heading("7. Zephyr Scale — Test Plans", level=1)
    doc.add_paragraph(
        "Test Plans provide high-level test planning that groups test cycles together. "
        "A test plan defines the scope and objectives of testing for a release or milestone, "
        "and can include multiple test cycles organized by feature area or testing phase."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-test-plans.png"),
        "Figure 12: Zephyr Scale — Test Plans View (Test Planning Interface)",
    )

    # ── Section 8: Reports ──
    doc.add_page_break()
    doc.add_heading("8. Zephyr Scale — Reports & Dashboards", level=1)
    doc.add_paragraph(
        "Zephyr Scale provides built-in test reporting and dashboards for tracking "
        "test execution progress, test case coverage, and defect metrics. Reports "
        "include coverage analysis, execution status, and traceability matrices."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "zephyr-reports.png"),
        "Figure 13: Zephyr Scale — Reports Dashboard (Test Metrics & Analytics)",
    )

    # ── Section 9: Features Used ──
    doc.add_page_break()
    doc.add_heading("9. Zephyr Scale Features Used", level=1)
    doc.add_paragraph(
        "The following Zephyr Scale features were utilized in this assignment:"
    )
    features = [
        "Test Cases Tab — Centralized test case repository with search and filtering",
        "Test Case Versioning — Each test case has version tracking (Version 1.0)",
        "Step-by-Step Test Scripts — Structured test steps with Test Data and Expected Results",
        "Folder Organization — Test cases organized into folders by module",
        "Status Workflow — Draft → Approved → Deprecated lifecycle management",
        "Priority Levels — Normal, High, Low, Critical for test prioritization",
        "Owner Assignment — Test cases assigned to specific team members",
        "Traceability — Link test cases to Jira issues for requirement coverage",
        "Test Cycles — Group test cases together for sprint-based execution",
        "Test Plans — High-level test planning for releases and milestones",
        "Reports — Built-in test metrics, coverage analysis, and dashboards",
    ]
    for feat in features:
        doc.add_paragraph(feat, style="List Bullet")

    # ── Section 10: Comparison ──
    doc.add_page_break()
    doc.add_heading("10. Jira Issues vs Zephyr Scale — Comparison", level=1)
    doc.add_paragraph(
        "The table below compares using plain Jira issues versus Zephyr Scale "
        "for test case management:"
    )

    comp_table = doc.add_table(rows=1, cols=3)
    comp_table.style = "Light Grid Accent 1"
    comp_hdr = comp_table.rows[0].cells
    comp_hdr[0].text = "Feature"
    comp_hdr[1].text = "Jira Issues"
    comp_hdr[2].text = "Zephyr Scale"
    comparisons = [
        ("Test Steps", "Manual (in description)", "Structured step editor"),
        ("Expected Results", "Text in description", "Per-step expected results"),
        ("Test Execution", "Manual status change", "Dedicated execution engine"),
        ("Test Cycles", "Not available", "Built-in test cycles"),
        ("Test Plans", "Epics/Stories", "Dedicated test plans"),
        ("Reporting", "Jira dashboards", "Test-specific reports"),
        ("Versioning", "Not available", "Test case versioning"),
        ("Traceability", "Links only", "Requirement coverage matrix"),
    ]
    for feature, jira, zephyr in comparisons:
        row = comp_table.add_row().cells
        row[0].text = feature
        row[1].text = jira
        row[2].text = zephyr
    doc.add_paragraph()

    # ── Section 11: Jira Project — Supporting Evidence ──
    doc.add_page_break()
    doc.add_heading("11. Jira Project — Supporting Evidence", level=1)
    doc.add_paragraph(
        "The following screenshots provide supporting evidence of the MegiLance "
        "Jira project setup, demonstrating the full Agile/Scrum workflow alongside "
        "the Zephyr Scale test management. The project contains 88+ work items across "
        "multiple issue types with sprint planning, automation rules, and timeline tracking."
    )

    doc.add_heading("11.1 Scrum Board", level=2)
    doc.add_paragraph(
        "The MegiLance Scrum board shows work items organized in TO DO, IN PROGRESS, "
        "and DONE columns for sprint-based development tracking."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "03-scrum-board.png"),
        "Figure 14: MEG Scrum Board — TO DO / IN PROGRESS / DONE columns",
    )

    doc.add_heading("11.2 Backlog & Sprint Planning", level=2)
    doc.add_paragraph(
        "Sprints are organized in the backlog with defined start/end dates and assigned "
        "work items, demonstrating proper Scrum sprint planning methodology."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "04-backlog-sprints.png"),
        "Figure 15: Backlog — Sprint planning with MEG Sprint 1 and Foundation Sprint",
    )

    doc.add_heading("11.3 Timeline / Epics", level=2)
    doc.add_paragraph(
        "Epics provide high-level feature grouping across the project timeline, "
        "including Authentication, AI Matching, Client/Freelancer/Admin Portals."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "05-timeline-epics.png"),
        "Figure 16: Timeline view — Epics (Auth, AI Matching, Client/Freelancer/Admin Portals)",
    )

    doc.add_heading("11.4 All Work Items (88 Issues)", level=2)
    doc.add_paragraph(
        "The project tracks 88+ issues across multiple types: Epics, Stories, Bugs, "
        "and Tasks — spanning all modules of the MegiLance platform."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "07-all-issues-types.png"),
        "Figure 17: All work items — 88+ issues filtered by type (Epic, Story, Bug, Task)",
    )

    doc.add_heading("11.5 Automation Rules", level=2)
    doc.add_paragraph(
        "Jira automation rules streamline the workflow with automated actions for "
        "bug priority-based due dates, sub-task completion rollups, due date reminders, "
        "and auto-assignment on transitions."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "06-automation-rules.png"),
        "Figure 18: All automation rules — enabled and active for workflow automation",
    )

    doc.add_heading("11.6 Jira Test Case Issues (21 Issues)", level=2)
    doc.add_paragraph(
        "In addition to the 10 Zephyr Scale test cases, 21 test case issues were created "
        "directly in Jira with the 'test-case' label, providing comprehensive test coverage "
        "documentation within the standard Jira issue tracker."
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "01-test-cases-list.png"),
        "Figure 19: Jira Issues — 21 Test Cases with label 'test-case' (JQL filtered)",
    )

    # Show sample Jira issue details
    doc.add_paragraph(
        "Sample Jira test case issue details showing structured test steps, "
        "preconditions, expected results, and labels:"
    )
    add_screenshot(
        doc,
        os.path.join(SCREENSHOTS_DIR, "02-test-case-detail.png"),
        "Figure 20: MEG-105 — TC-PROJ-001: Create New Project with Valid Data (Jira Issue Detail)",
    )

    # ── Section 12: Conclusion ──
    doc.add_page_break()
    doc.add_heading("12. Conclusion", level=1)
    doc.add_paragraph(
        "Zephyr Scale provides a significantly more structured and professional approach "
        "to test case management compared to plain Jira issues. The dedicated test case "
        "editor, step-by-step test scripts, execution tracking, and reporting capabilities "
        "make it the preferred tool for managing test cases in real-world software testing projects."
    )
    doc.add_paragraph(
        "For the MegiLance FYP project, 10 test cases were successfully created in Zephyr Scale "
        "covering authentication, project management, user profiles, proposals, payments, and "
        "search modules — demonstrating the tool's capabilities for structured test management."
    )
    doc.add_paragraph(
        "Key achievements demonstrated in this assignment:"
    )
    achievements = [
        "10 structured test cases created in Zephyr Scale (MEG-T1 to MEG-T10)",
        "21 test case issues in Jira with proper labels and categorization",
        "88+ total work items across Epics, Stories, Bugs, and Tasks",
        "Step-by-step test scripts with Test Data and Expected Results",
        "Test Cycles and Test Plans configured for execution tracking",
        "Jira automation rules for workflow efficiency",
        "Full Scrum board with sprint planning and backlog management",
        "Timeline/Epic-level project tracking across all modules",
    ]
    for a in achievements:
        doc.add_paragraph(a, style="List Bullet")

    # ── Section 13: URLs Reference ──
    add_url_section(doc, "13. URLs Reference", [
        ("Zephyr Scale — All Test Cases",
         "https://ghulam-mujtaba.atlassian.net/projects/MEG?selectedItem=com.atlassian.plugins.atlassian-connect-plugin:com.kanoah.test-manager__main-project-page"),
        ("Zephyr Scale — Atlassian Marketplace",
         "https://marketplace.atlassian.com/apps/1213259/zephyr-test-management-and-automation-for-jira"),
        ("Jira Project — MegiLance (MEG) — Scrum Board",
         "https://ghulam-mujtaba.atlassian.net/jira/software/c/projects/MEG/boards/3"),
        ("Backlog & Sprint Planning",
         "https://ghulam-mujtaba.atlassian.net/jira/software/c/projects/MEG/boards/3/backlog"),
        ("Timeline / Epics",
         "https://ghulam-mujtaba.atlassian.net/jira/software/c/projects/MEG/boards/3/timeline"),
        ("All 21 Jira Test Case Issues",
         "https://ghulam-mujtaba.atlassian.net/issues/?jql=project%20%3D%20MEG%20AND%20labels%20%3D%20test-case"),
        ("All Issues (Epic, Story, Bug, Task)",
         "https://ghulam-mujtaba.atlassian.net/issues/?jql=project%20%3D%20MEG%20AND%20type%20in%20(Epic%2C%20Story%2C%20Bug%2C%20Task)%20ORDER%20BY%20type%20ASC"),
        ("Automation Rules",
         "https://ghulam-mujtaba.atlassian.net/jira/software/c/projects/MEG/settings/automate"),
        ("TC-PROJ-001 — Create New Project (MEG-105)",
         "https://ghulam-mujtaba.atlassian.net/browse/MEG-105"),
    ])

    output_path = os.path.join(OUTPUT_DIR, "CSE455_ZephyrScale_TestCases_GhulamMujtaba_v2.docx")
    doc.save(output_path)
    print(f"Zephyr Scale Assignment DOCX saved: {output_path}")


if __name__ == "__main__":
    create_assignment1()
    create_assignment2()
    create_zephyr_assignment()
    print("\nDone! All DOCX files created.")
